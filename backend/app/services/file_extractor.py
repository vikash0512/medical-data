from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import List

from .exceptions import ExtractionError


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


@dataclass
class FileExtractionResult:
    title: str
    blocks: List[str]
    source_name: str


def split_text_blocks(text: str) -> List[str]:
    candidates = [chunk.strip() for chunk in text.replace("\r\n", "\n").split("\n\n")]
    if len(candidates) <= 1:
        candidates = [line.strip() for line in text.splitlines()]
    return [candidate for candidate in candidates if candidate]


def extract_txt(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractionError("TXT file could not be decoded as UTF-8, UTF-16, or Latin-1.")


def extract_pdf(content: bytes) -> str:
    try:
        from pdfminer.high_level import extract_text

        return extract_text(BytesIO(content))
    except ImportError:
        try:
            import fitz
        except ImportError as exc:
            raise ExtractionError(
                "PDF extraction requires pdfminer.six or PyMuPDF. Install backend requirements first."
            ) from exc

        document = fitz.open(stream=content, filetype="pdf")
        return "\n\n".join(page.get_text() for page in document)


def extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ExtractionError(
            "DOCX extraction requires python-docx. Install backend requirements first."
        ) from exc

    document = Document(BytesIO(content))
    parts: List[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


def extract_file(filename: str, content: bytes) -> FileExtractionResult:
    if not filename:
        raise ExtractionError("Uploaded file is missing a filename.")

    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ExtractionError(f"Unsupported file type. Supported types: {supported}.")

    if extension == ".txt":
        text = extract_txt(content)
    elif extension == ".pdf":
        text = extract_pdf(content)
    else:
        text = extract_docx(content)

    blocks = split_text_blocks(text)
    if not blocks:
        raise ExtractionError("No readable text was found in the uploaded file.")

    return FileExtractionResult(title=Path(filename).stem, blocks=blocks, source_name=filename)

